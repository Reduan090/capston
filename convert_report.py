#!/usr/bin/env python3
"""Convert TECHNICAL_REPORT.md to Word and PDF formats"""

import subprocess
from pathlib import Path

report_path = Path("TECHNICAL_REPORT.md")
output_docx = Path("TECHNICAL_REPORT.docx")
output_pdf = Path("TECHNICAL_REPORT.pdf")

print(f"Converting {report_path} to DOCX and PDF...")

try:
    # Install pandoc if needed
    subprocess.run(["pandoc", "--version"], capture_output=True, check=True)
except Exception as e:
    print(f"⚠️  Pandoc not installed. Install from: https://pandoc.org/installing.html")
    print(f"Error: {e}")
    print("\nAttempting alternative conversion using python-docx...")

# Try conversion with pandoc
try:
    # Convert to DOCX
    cmd_docx = ["pandoc", str(report_path), "-o", str(output_docx), 
                "--from", "markdown", "--to", "docx"]
    result = subprocess.run(cmd_docx, capture_output=True, text=True)
    if result.returncode == 0:
        print(f"✅ Created: {output_docx}")
    else:
        print(f"❌ DOCX conversion failed: {result.stderr}")
    
    # Convert to PDF (requires wkhtmltopdf or similar)
    cmd_pdf = ["pandoc", str(report_path), "-o", str(output_pdf),
               "--from", "markdown", "--to", "pdf"]
    result = subprocess.run(cmd_pdf, capture_output=True, text=True)
    if result.returncode == 0:
        print(f"✅ Created: {output_pdf}")
    else:
        print(f"⚠️  PDF conversion requires LaTeX/wkhtmltopdf: {result.stderr}")

except FileNotFoundError:
    print("❌ Pandoc not found. Install from: https://pandoc.org/installing.html")
    print("\nUsing alternative: Converting to DOCX with python-docx...")
    
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Read markdown and convert
        doc = Document()
        with open(report_path) as f:
            content = f.read()
        
        # Simple markdown parsing (for basic conversion)
        lines = content.split('\n')
        for line in lines:
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith('```'):
                continue
            elif line.strip():
                doc.add_paragraph(line)
        
        doc.save(output_docx)
        print(f"✅ Created: {output_docx} (basic formatting)")
    
    except ImportError:
        print("❌ python-docx not installed. Install: pip install python-docx")

print("\n📄 Report files ready!")
print(f"Markdown: {report_path}")
print(f"Word (DOCX): {output_docx}")
print(f"PDF: {output_pdf}")
