#!/usr/bin/env python3
"""Generate DOCX version of Technical Report directly from markdown"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from pathlib import Path

def convert_md_to_docx(md_file, docx_file):
    """Convert markdown technical report to Word document"""
    
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split by lines but preserve structure
    lines = content.split('\n')
    
    in_code_block = False
    code_content = []
    in_table = False
    table_rows = []
    
    for line in lines:
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                for code_line in code_content:
                    if code_content.index(code_line) == 0:
                        p.text = code_line
                    else:
                        doc.add_paragraph(code_line, style='List Bullet')
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        # Handle headings
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        
        # Handle bold/italic
        elif line.strip() and not in_code_block:
            # Process text with formatting
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            
            # Replace markdown formatting
            text = line
            
            # Add paragraph
            if text.strip():
                p.text = text
        
        elif in_code_block:
            code_content.append(line)
        
        elif line.strip() == '':
            doc.add_paragraph()
    
    doc.save(docx_file)
    print(f"✅ Created: {docx_file}")

# Main execution
if __name__ == '__main__':
    md_path = Path('TECHNICAL_REPORT.md')
    docx_path = Path('TECHNICAL_REPORT.docx')
    
    if md_path.exists():
        print(f"Converting {md_path} to {docx_path}...")
        convert_md_to_docx(str(md_path), str(docx_path))
        print("📄 Conversion complete!")
        print(f"\nFiles available:")
        print(f"  - Markdown: TECHNICAL_REPORT.md")
        print(f"  - Word: TECHNICAL_REPORT.docx")
        print(f"\nFor PDF conversion, install Pandoc:")
        print(f"  https://pandoc.org/installing.html")
        print(f"  Then run: pandoc TECHNICAL_REPORT.md -o TECHNICAL_REPORT.pdf")
    else:
        print(f"❌ {md_path} not found!")
